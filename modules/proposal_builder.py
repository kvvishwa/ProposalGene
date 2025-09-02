# modules/proposal_builder.py
# -----------------------------------------------------------------------------
# Build a proposal DOCX from a template, static sections, and dynamic sections.
# Ordering: STATIC sections first (via XML-preserving insertion), then DYNAMIC
# sections (each with a proper heading). Uses your docx_assembler/docx_generator
# helpers to preserve formatting, lists, tables, and styles.
# -----------------------------------------------------------------------------

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from copy import deepcopy
from typing import Dict, List, Optional, Tuple, Any

# python-docx
from docx import Document
from docx.document import Document as _Doc
from docx.text.paragraph import Paragraph

# ---- Rich DOCX helpers (preserve formatting) --------------------------------
# insert_section: anchor-aware, XML-level merge
try:
    from modules.docx_assembler import insert_section as _insert_section_docx
except Exception:
    from docx_assembler import insert_section as _insert_section_docx  # type: ignore

# style-aware heading, markdown, sources, toc, static heading normalization
try:
    from modules.docx_generator import (
        add_page_break as _add_page_break,
        add_heading_paragraph as _add_heading_paragraph,
        insert_markdown_block as _insert_markdown_block,
        add_sources_line as _add_sources_line,
        insert_toc as _insert_toc,
        normalize_static_section_heading as _normalize_static_section_heading,
    )
except Exception:
    from modules.docx_generator import (  # type: ignore
        add_page_break as _add_page_break,
        add_heading_paragraph as _add_heading_paragraph,
        insert_markdown_block as _insert_markdown_block,
        add_sources_line as _add_sources_line,
        insert_toc as _insert_toc,
        normalize_static_section_heading as _normalize_static_section_heading,
    )

# ---- Evidence retrieval ------------------------------------------------------
try:
    from modules.sp_retrieval import get_sp_evidence_for_question as _get_sp_evidence_q
except Exception:
    try:
        from sp_retrieval import get_sp_evidence_for_question as _get_sp_evidence_q  # type: ignore
    except Exception:
        _get_sp_evidence_q = None  # fallback to empty evidence


# =============================================================================
# Build options
# =============================================================================
@dataclass
class BuildOptions:
    # General
    use_anchors: bool = True                 # static: use template anchors if present
    template_has_headings: bool = True
    page_breaks: bool = True
    include_sources: bool = True
    add_toc: bool = True

    # Dynamic composition (style)
    rec_style: str = "bullets"               # "bullets" | "paragraphs"
    top_k_default: int = 6
    top_k_per_section: int = 0               # 0 → use default
    tone: str = "Professional"

    # Static heading normalization for inserted static .docx
    # 'keep'   -> leave first heading untouched
    # 'demote' -> Heading N -> Heading N+1 (avoids clashing with template H1)
    # 'strip'  -> remove heading text in the inserted section's first paragraph
    static_heading_mode: str = "demote"

    # Dynamic headings
    dynamic_heading_level: int = 1
    dynamic_heading_style_name: Optional[str] = None
    force_dynamic_heading: bool = True       # ALWAYS add a heading for dynamic sections

    # Facts/context (optional; used by retrieval prompts)
    facts: dict = field(default_factory=dict)


# =============================================================================
# Evidence retrieval + dynamic composition
# =============================================================================
def _get_evidence(label: str, facts: dict, cfg, k: int) -> List[Any]:
    if _get_sp_evidence_q is None:
        return []
    try:
        return _get_sp_evidence_q(label, facts or None, cfg, k=k) or []
    except Exception:
        return []


def _compose_dynamic_markdown(label: str, ev: List[Any], oai, cfg, opts: BuildOptions) -> Tuple[str, List[dict]]:
    """
    Turn evidence into a numbered context and ask the LLM to write the section.
    Returns (markdown_text, normalized_evidence_dicts).
    """
    if not ev:
        return (
            f"*No high-confidence SharePoint evidence found for **{label}**. "
            f"Consider re-indexing or adjusting your query.*",
            []
        )

    numbered = []
    norm_ev: List[dict] = []
    seen_srcs = set()

    for i, e in enumerate(ev, start=1):
        if isinstance(e, dict):
            src = e.get("source")
            pg  = e.get("page_hint")
            txt = e.get("text", "")
            norm_ev.append({"source": src, "page_hint": pg, "text": txt})
        else:
            src = getattr(e, "source", None)
            pg  = getattr(e, "page_hint", None)
            txt = getattr(e, "text", "")
            norm_ev.append({"source": src, "page_hint": pg, "text": txt})

        src_str = f"{src}{(' p.' + str(pg)) if pg else ''}" if src else f"[evidence {i}]"
        if src_str not in seen_srcs:
            seen_srcs.add(src_str)
        numbered.append(f"[{i}] {src_str}\n{txt}")

    context = "\n\n".join(numbered)
    style_hint = "bullet points" if (opts.rec_style or "").lower().startswith("bullet") else "tight paragraphs"
    prompt = (
        f"Write the section **{label}** for a proposal.\n"
        f"- style: {style_hint}, tone: {opts.tone}\n"
        "- Base it strictly on the numbered evidence; do not add external facts.\n"
        "- Append [^n] to claims grounded in evidence [n].\n\n"
        f"EVIDENCE:\n{context}\n"
    )

    try:
        res = oai.chat.completions.create(
            model=getattr(cfg, "ANALYSIS_MODEL", "gpt-4o"),
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        md = (res.choices[0].message.content or "").strip()
    except Exception as ex:
        md = f"*LLM error while composing **{label}***: {ex}"

    return md, norm_ev


# =============================================================================
# Static insertion (formatting-preserving)
# =============================================================================
def _insert_static_section_preserve(base_doc: _Doc, label: str, path: str, opts: BuildOptions) -> None:
    """
    Insert a static DOCX into base_doc using XML-preserving merge:
      - Try SDT/tag/alias anchor or bookmark; else append.
      - Optionally demote/strip first heading INSIDE the section doc
        (to avoid clashing with template headings).
    """
    section_doc = Document(path)
    _normalize_static_section_heading(section_doc, mode=opts.static_heading_mode)

    # When appending (no anchor), we ask insert_section() to add a Heading 1
    # with the section label. If your template already has a heading/anchor, it won't add.
    try:
        _insert_section_docx(
            base_doc,
            section_label=label,
            section_doc=section_doc,
            try_anchors=opts.use_anchors,
            add_heading_if_appending=True,
        )
    except Exception as ex:
        # Hard fallback: append a page break + simple heading + copy paragraphs (last resort)
        if opts.page_breaks:
            _add_page_break(base_doc)
        _add_heading_paragraph(base_doc, label)
        for p in section_doc.paragraphs:
            newp = base_doc.add_paragraph()
            for r in p.runs:
                run = newp.add_run(r.text)
                try:
                    run.bold = r.bold
                    run.italic = r.italic
                    run.underline = r.underline
                except Exception:
                    pass


# =============================================================================
# Dynamic insertion (always with heading)
# =============================================================================
def _insert_dynamic_section(base_doc: _Doc, label: str, md: str, evidence: List[dict], opts: BuildOptions) -> None:
    """
    Append a dynamic section:
      - Optional page break
      - Styled heading (level/style from opts)
      - Markdown-ish body
      - Optional 'Sources:' line
    """
    if opts.page_breaks:
        _add_page_break(base_doc)
    _add_heading_paragraph(
        base_doc,
        label,
        style_name=(opts.dynamic_heading_style_name or f"Heading {max(1, min(9, opts.dynamic_heading_level))}"),
    )
    last = _insert_markdown_block(base_doc, md)
    if opts.include_sources and evidence:
        srcs = list({e.get("source") for e in evidence if e.get("source")})
        if srcs:
            _add_sources_line(base_doc, srcs, after_paragraph=last)


# =============================================================================
# Main builder
# =============================================================================
def build_proposal(
    template_path: str,
    static_paths: Dict[str, str],     # label -> path for static DOCX sections
    final_order: List[str],           # mixed labels; convention: dynamics prefixed with "[Dyn] "
    oai,
    cfg,
    opts: BuildOptions,
):
    """
    Build the final proposal DOCX + optional recommendations/preview DOCX.

    Returns:
        draft_docx_bytes: bytes
        recos_docx_bytes: Optional[bytes]
        preview_dict: Dict[str, Any]  (per dynamic section: md, sources, evidence)
        meta: Dict[str, Any]
    """
    base_doc = Document(template_path)

    # Optional TOC at top
    if opts.add_toc:
        try:
            _insert_toc(base_doc, title="Table of Contents")
        except Exception:
            pass

    # ---------- enforce STATIC-FIRST ordering ----------
    static_labels  = [lbl for lbl in final_order if not lbl.startswith("[Dyn] ") and lbl in static_paths]
    dynamic_items  = [lbl for lbl in final_order if lbl.startswith("[Dyn] ")]
    final_order2   = static_labels + dynamic_items

    preview: Dict[str, dict] = {}

    # ---------- 1) insert all STATIC sections (preserve formatting) ----------
    for label in static_labels:
        path = static_paths.get(label)
        if not path:
            continue
        _insert_static_section_preserve(base_doc, label, path, opts)

    # ---------- 2) insert all DYNAMIC sections (append + heading) ----------
    for dyn in dynamic_items:
        label = dyn[6:].strip()  # strip "[Dyn] "
        k = opts.top_k_per_section or opts.top_k_default
        ev = _get_evidence(label, opts.facts, cfg, k=k)
        md, ev_norm = _compose_dynamic_markdown(label, ev, oai, cfg, opts)
        _insert_dynamic_section(base_doc, label, md, ev_norm, opts)

        preview[label] = {
            "md": md,
            "sources": list({e.get("source") for e in ev_norm if e.get("source")}),
            "evidence": ev_norm,
        }

    # ---------- Save main draft ----------
    b1 = BytesIO()
    base_doc.save(b1)
    draft_bytes = b1.getvalue()

    # ---------- Optional: build a "Dynamic Recommendations" preview DOCX ----------
    recs_bytes: Optional[bytes] = None
    if preview:
        doc2 = Document()
        _add_heading_paragraph(doc2, "Dynamic Recommendations (Preview)", style_name="Heading 1")
        for sec, pkg in preview.items():
            _add_heading_paragraph(doc2, sec, style_name="Heading 2")
            last = _insert_markdown_block(doc2, pkg["md"])
            if (opts.include_sources and pkg.get("sources")):
                _add_sources_line(doc2, pkg["sources"], after_paragraph=last)
        b2 = BytesIO()
        doc2.save(b2)
        recs_bytes = b2.getvalue()

    meta = {"template": template_path, "final_order": final_order2}
    return draft_bytes, recs_bytes, preview, meta
