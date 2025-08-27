import os
import tempfile

def save_to_temp(uploaded_file):
    suffix = uploaded_file.name.split('.')[-1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{suffix}') as tmp:
        tmp.write(uploaded_file.getvalue())
        return tmp.name

def cleanup_temp_files(paths):
    if not paths:
        return
    for p in list(paths):
        try:
            if p and os.path.exists(p) and os.path.isfile(p):
                os.remove(p)
        except Exception:
            pass


# modules/utils.py  (add this)
from io import BytesIO
from typing import Union
from docx import Document as _DocxDocument

def safe_read_docx(path_or_bytes: Union[str, bytes]) -> _DocxDocument:
    """
    Robust loader for .docx that works with a filesystem path or raw bytes.
    Falls back to a BytesIO stream if direct open fails.
    """
    try:
        # If it's bytes already
        if isinstance(path_or_bytes, (bytes, bytearray)):
            return _DocxDocument(BytesIO(path_or_bytes))
        # Else assume it's a path-like
        return _DocxDocument(path_or_bytes)
    except Exception:
        # Last-ditch: open file and feed bytes
        if isinstance(path_or_bytes, (str,)):
            with open(path_or_bytes, "rb") as f:
                return _DocxDocument(BytesIO(f.read()))
        raise
