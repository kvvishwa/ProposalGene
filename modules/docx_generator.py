# modules/docx_generator.py
# -----------------------------------------------------------------------------
# Word (.docx) helpers used by the proposal builder.
# - Safe insertion after a given paragraph (via OXML)
# - Style-aware heading insertion
# - Append another docx after a specific paragraph index
# - Minimal Markdown → Word block
# - Sources line & page break helpers
# - Anchor finder (lightweight)
# - Static heading normalization
# -----------------------------------------------------------------------------

from __future__ import annotations

from typing import List, Optional
from copy import deepcopy

from docx import Document
from docx.document import Document as _Doc
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.shared import OxmlElement, qn

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------- insertion primitives -----------------------------

# --- add near the top ---
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import re
_XML_ILLEGAL = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')
def sanitize_for_xml(text: str) -> str:
    return _XML_ILLEGAL.sub('', text or '')

def _local(tag):
    return tag.split('}', 1)[1] if '}' in tag else tag

def _strip_numPr_in_paragraph(p_elm):
    """Remove numbering from a paragraph to avoid missing num/abstractNum definitions."""
    if _local(p_elm.tag) != "p":
        return
    pPr = None
    for ch in p_elm:
        if _local(ch.tag) == "pPr":
            pPr = ch
            break
    if pPr is None:
        return
    to_remove = []
    for ch in pPr:
        if _local(ch.tag) == "numPr":
            to_remove.append(ch)
    for n in to_remove:
        pPr.remove(n)

def _strip_nested_numPr(elm):
    """Walk element and remove numPr in any descendant paragraphs (incl. tables)."""
    if _local(elm.tag) == "p":
        _strip_numPr_in_paragraph(elm)
    for ch in list(elm):
        _strip_nested_numPr(ch)

def _is_sectPr(elm):
    return _local(elm.tag) == "sectPr"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: Optional[str] = None) -> Paragraph:
    """
    Insert a new paragraph node *immediately after* the given paragraph using OXML.
    Returns the new Paragraph.
    """
    new_p_elm = OxmlElement('w:p')
    paragraph._p.addnext(new_p_elm)
    new_p = Paragraph(new_p_elm, paragraph._parent)
    if style_name:
        try:
            new_p.style = style_name
        except Exception:
            pass
    if text:
        new_p.add_run(text)
    return new_p


def add_page_break(doc: _Doc, after_paragraph: Optional[Paragraph] = None) -> Paragraph:
    """
    Insert a page break. If after_paragraph is provided, insert right after it,
    else append at end of the document.
    Returns the paragraph containing the break.
    """
    if after_paragraph is not None:
        p = insert_paragraph_after(after_paragraph)
    else:
        p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def add_heading_paragraph(doc: _Doc, text: str, style_name: str = "Heading 1", after_paragraph: Optional[Paragraph] = None) -> Paragraph:
    """
    Insert a heading paragraph, using the provided style name (falls back gracefully).
    If after_paragraph is provided, insert right after it.
    """
    if after_paragraph is not None:
        p = insert_paragraph_after(after_paragraph, text="", style_name=style_name)
    else:
        p = doc.add_paragraph()
        try:
            p.style = style_name
        except Exception:
            try:
                p.style = "Heading 1"
            except Exception:
                pass
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if text:
        p.add_run(text)
    return p


# ----------------------------- block builders -----------------------------

def insert_markdown_block(doc: _Doc, md_text: str, after_paragraph: Optional[Paragraph] = None) -> Paragraph:
    """
    Extremely small Markdown → Word helper.
    Supports:
      - bullet lines: starting with -, *, •
      - paragraphs otherwise
    Inserts content *after* the given paragraph if provided, keeping order.
    Returns the last paragraph inserted (useful to place 'Sources:' after).
    """
    lines = [ln.rstrip() for ln in (md_text or "").splitlines()]
    anchor = after_paragraph
    last_p = anchor

    def _insert_para(text: str, bullet: bool = False) -> Paragraph:
        nonlocal last_p
        if last_p is None:
            p = doc.add_paragraph()
        else:
            p = insert_paragraph_after(last_p)
        if bullet:
            # Try to use List Paragraph if available
            try:
                p.style = doc.styles.get("List Paragraph") or doc.styles["Normal"]
            except Exception:
                pass
        if text:
            p.add_run(text)
        last_p = p
        return p

    for ln in lines:
        if not ln.strip():
            _insert_para("")  # blank line
            continue
        if ln.lstrip().startswith(("-", "*", "•")):
            _insert_para(ln.lstrip("-*• ").strip(), bullet=True)
        else:
            _insert_para(ln, bullet=False)

    return last_p if last_p is not None else doc.paragraphs[-1]


def add_sources_line(doc: _Doc, sources: List[str], after_paragraph: Optional[Paragraph] = None) -> Paragraph:
    """
    Append a small 'Sources:' line, optionally inserted after a given paragraph.
    Returns the paragraph created.
    """
    text = "Sources: " + ", ".join(sources)
    if after_paragraph is None:
        p = doc.add_paragraph()
        p.add_run(text)
        return p
    else:
        p = insert_paragraph_after(after_paragraph)
        p.add_run(text)
        return p


def append_docx(base_doc: _Doc, add_doc: _Doc, after_index: Optional[int]) -> None:
    """
    Append contents of add_doc into base_doc.
    - Strips numbering (w:numPr) from inserted paragraphs.
    - Skips section properties (w:sectPr) blocks.
    - Works when after_index is provided (insert immediately after that paragraph).
    """
    body = base_doc.element.body
    add_body_elems = [el for el in add_doc.element.body if not _is_sectPr(el)]

    # sanitize numbering to avoid unreadable content
    safe_elems = []
    for el in add_body_elems:
        el2 = deepcopy(el)
        _strip_nested_numPr(el2)
        safe_elems.append(el2)

    if after_index is None:
        for el in safe_elems:
            body.append(el)
        return

    # Find insertion position by paragraph index
    try:
        anchor_p = base_doc.paragraphs[after_index]
        body_elems = list(body)
        pos = body_elems.index(anchor_p._p)
    except Exception:
        # Fallback to append at end
        for el in safe_elems:
            body.append(el)
        return

    insert_pos = pos + 1
    for el in safe_elems:
        body.insert(insert_pos, el)
        insert_pos += 1



# ----------------------------- anchors & normalization -----------------------------

def find_anchor_paragraph(doc: _Doc, label: str) -> Optional[Paragraph]:
    """
    Try to locate an anchor by (priority):
      1) Content control (SDT) paragraph whose text includes [[label]] or {{label}}
      2) A paragraph whose text exactly equals the label
    (Bookmark walking is omitted here for brevity.)
    Returns the Paragraph if found, else None.
    """
    needle1 = f"[[{label}]]".lower()
    needle2 = f"{{{{{label}}}}}".lower()
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        tl = t.lower()
        if needle1 in tl or needle2 in tl:
            return p
        if tl == label.lower():
            return p
    return None


def normalize_static_section_heading(doc: _Doc, mode: str = "demote") -> None:
    """
    Normalize first heading in a static section doc:
      - 'keep': leave untouched
      - 'demote': if first paragraph is Heading 1/2, demote to next level to avoid clashing with template H1
      - 'strip': if first paragraph looks like a heading, remove its text content
    """
    if mode not in ("keep", "demote", "strip"):
        mode = "demote"

    if not doc.paragraphs:
        return

    p0 = doc.paragraphs[0]
    name = (p0.style.name if p0.style else "") or ""
    is_heading = name.lower().startswith("heading")
    if not is_heading:
        return

    if mode == "strip":
        # Remove text runs (leave blank)
        for r in list(p0.runs):
            r.text = ""
        return

    if mode == "demote":
        try:
            if name.startswith("Heading "):
                lvl = int(name.split(" ")[1])
                new_name = f"Heading {min(6, lvl + 1)}"
                p0.style = doc.styles.get(new_name, p0.style)
        except Exception:
            pass


# --- Compat shim: insert_toc (so old imports don't break) ---
def insert_toc(doc, title: str = "Table of Contents", heading_level: int = 1):
    """
    Minimal TOC field. Word will populate it on first open (Update Field/F9).
    Keeps the function name expected by older code paths.
    """
    try:
        # Prefer a heading style; fall back gracefully
        try:
            style_name = f"Heading {heading_level}"
        except Exception:
            style_name = "Heading 1"

        # use the existing helper in this module
        add_heading_paragraph(doc, title, style_name=style_name)

        # Insert TOC field
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = doc.add_paragraph()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        p._p.append(fld)
    except Exception:
        # Fail-closed: just add a plain heading if field insertion fails
        try:
            add_heading_paragraph(doc, title, style_name="Heading 1")
        except Exception:
            doc.add_paragraph(title)


def insert_table_of_contents(doc, title: str = "Table of Contents", heading_level: int = 1):
    return insert_toc(doc, title=title, heading_level=heading_level)