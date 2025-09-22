# modules/app_helpers.py
# -----------------------------------------------------------------------------
# Shared helpers for:
# - FS & simple persistence (urls, static map, ingested map, blueprints, facts json)
# - Streamlit rerun compat
# - Vector search bridges (SharePoint & Uploaded stores) + simple RAG answerer
# - TOC insertion into DOCX
# - Small utilities for structuring Q&A filtering
# -----------------------------------------------------------------------------

from __future__ import annotations
import os
import json
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
# add near the top
import re
import streamlit as st

# Vector stores (your implementations)
from modules.vectorstore import (
    init_sharepoint_store,
    init_uploaded_store,
)

# ---------------------------------------------------------------------
# Filesystem & persistence
# ---------------------------------------------------------------------

def _derive_cues_from_prompt(prompt: str) -> List[str]:
    """
    If the prompt includes 'Context focus cues:' (your sectional prompts do),
    use that line as high-signal retrieval seeds. Otherwise return [].
    """
    if not prompt:
        return []
    m = re.search(r'^\s*Context\s+focus\s+cues:\s*(.+)$', prompt, flags=re.I | re.M)
    if not m:
        return []
    # split on commas/semicolons, keep short phrases
    raw = m.group(1)
    cues = [c.strip() for c in re.split(r'[;,]', raw) if c.strip()]
    return cues[:6]

def _gather_understanding_context(store: Any, prompt: str, base_k: int = 8) -> Tuple[str, List[str]]:
    """
    Build a big, diverse context for the extractor using multiple compact cues.
    Uses the same tolerant cascade used in SharePoint chat.
    """
    try:
        try:
            from modules.sp_retrieval import get_store_evidence
        except ModuleNotFoundError:
            from sp_retrieval import get_store_evidence  # type: ignore
    except Exception:
        get_store_evidence = None

    # 1) cues from prompt (sectional calls provide these)
    cues = _derive_cues_from_prompt(prompt)

    # 2) standard understanding cues (cover all sections)
    std = [
        "contact email phone procurement communications",
        "submission instructions portal copies labeling registration format",
        "schedule due date deadlines pre-bid site visit Q&A addendum award",
        "evaluation criteria weighting selection process BAFO negotiations protest",
        "minimum qualifications insurance bonding certification years experience",
        "proposal organization page limits required sections forms pricing",
        "contract terms compliance privacy security scope of work",
    ]
    # prepend prompt-derived cues if any, then fallback to std list
    queries = cues + std

    pool = []
    seen = set()
    for q in queries:
        evs = []
        if get_store_evidence is not None:
            try:
                evs = get_store_evidence(store, question=q, rfp_facts=None, k=max(8, base_k))
            except Exception:
                evs = []
        # light fallback if cascade not available
        if not evs:
            try:
                docs = store.similarity_search(q, k=max(6, base_k))
            except Exception:
                docs = []
            for d in docs or []:
                txt = getattr(d, "page_content", "") or getattr(d, "content", "") or ""
                meta = getattr(d, "metadata", {}) or {}
                src  = meta.get("source", "") or meta.get("file", "") or ""
                if not txt.strip():
                    continue
                sig = (src + "|" + txt[:200]).lower()
                if sig in seen:
                    continue
                seen.add(sig)
                pool.append(type("E", (), {"text": txt[:1800], "source": src, "page_hint": meta.get("page")}))  # simple shim
                continue

        for e in evs or []:
            sig = ((getattr(e, "source", "") or "") + "|" + ((getattr(e, "text", "") or "")[:200])).lower()
            if sig in seen or not (getattr(e, "text", "") or "").strip():
                continue
            seen.add(sig)
            pool.append(e)

        if len(pool) >= 60:  # enough context
            break

    if not pool:
        return "", []

    # Build joined context & sources
    context = "\n\n---\n\n".join((getattr(e, "text", "") or "")[:1800] for e in pool[:60])
    sources = []
    for e in pool:
        src = getattr(e, "source", "") or ""
        if src and src not in sources:
            sources.append(src)
    return context, sources



def _derive_retrieval_query(prompt: str) -> str:
    """
    Prefer a short, high-signal query for vector retrieval.
    - If the prompt contains a leading 'Context focus cues:' line, use that line.
    - If the prompt is huge (schema prompts), fall back to broad RFP tokens.
    """
    if not prompt:
        return ""
    m = re.search(r'^\s*Context\s+focus\s+cues:\s*(.+)$', prompt, flags=re.I | re.M)
    if m:
        return m.group(1).strip()

    # If it's a giant schema prompt, use a generic but useful retrieval seed
    if len(prompt) > 800:
        return ("rfp procurement contact email phone submission instructions schedule deadline "
                "evaluation criteria compliance terms page limit mandatory forms")

    # Otherwise, use the first non-empty line (kept short)
    first_line = next((ln.strip() for ln in prompt.splitlines() if ln.strip()), "")
    return first_line[:512]


def ensure_dirs(base_dir: Path) -> None:
    """Create base subfolders used by the app (idempotent)."""
    (base_dir / "static_sources").mkdir(parents=True, exist_ok=True)
    (base_dir / "templates").mkdir(parents=True, exist_ok=True)
    (base_dir / "rfp_facts").mkdir(parents=True, exist_ok=True)

def list_files(folder: Path, exts: Tuple[str, ...] = ("docx",)) -> List[str]:
    """List files in a folder filtered by extensions (lowercase, no dots)."""
    if not folder.exists():
        return []
    out: List[str] = []
    extset = set(x.lower().lstrip(".") for x in exts)
    for p in sorted(folder.iterdir()):
        if p.is_file():
            ext = p.suffix.lower().lstrip(".")
            if ext in extset:
                out.append(p.name)
    return out

# ---- Saved SharePoint site URLs
def load_saved_urls(json_path: Path) -> List[str]:
    if not json_path.exists():
        return []
    try:
        return json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return []

def save_saved_urls(json_path: Path, urls: List[str]) -> None:
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(urls or [], ensure_ascii=False, indent=2), encoding="utf-8")

def upsert_url(urls: List[str], url: str, idx: Optional[int]) -> List[str]:
    """Insert or replace URL. If idx is None → append if not exists."""
    if idx is not None and 0 <= idx < len(urls):
        urls[idx] = url
    else:
        if url not in urls:
            urls.append(url)
    return urls

def delete_url(urls: List[str], idx: int) -> List[str]:
    if 0 <= idx < len(urls):
        del urls[idx]
    return urls

# ---- Static mapping (section -> filename)
def load_static_map(json_path: Path) -> Dict[str, str]:
    if not json_path.exists():
        return {}
    try:
        return json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return {}

def save_static_map(json_path: Path, mapping: Dict[str, str]) -> None:
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(mapping or {}, ensure_ascii=False, indent=2), encoding="utf-8")

# ---- SharePoint ingested files map (site_key -> [filepaths])
def load_sp_ingested_map(json_path: Path) -> Dict[str, List[str]]:
    if not json_path.exists():
        return {}
    try:
        return json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return {}

def save_sp_ingested_map(json_path: Path, mapping: Dict[str, List[str]]) -> None:
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(mapping or {}, ensure_ascii=False, indent=2), encoding="utf-8")

# ---- RFP facts JSON persistence
def save_rfp_facts_json(folder: Path, facts: dict, filename: str) -> Path:
    """Persist the extracted/edited facts JSON under /rfp_facts. Returns the saved path."""
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / filename
    with path.open("w", encoding="utf-8") as f:
        json.dump(facts or {}, f, ensure_ascii=False, indent=2)
    return path

def list_rfp_facts_json(folder: Path) -> List[str]:
    """List previously saved facts JSON files."""
    if not folder.exists():
        return []
    return sorted([p.name for p in folder.glob("*.json")])

def load_rfp_facts_json(folder: Path, filename: str) -> Optional[dict]:
    """Load a saved facts JSON by filename (in /rfp_facts)."""
    path = folder / filename
    if not path.exists():
        return None
    try:
        with path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None

# ---------------------------------------------------------------------
# Streamlit rerun compatibility
# ---------------------------------------------------------------------
def st_rerun_compat() -> None:
    """Call st.experimental_rerun if st.rerun not available (older Streamlit)."""
    try:
        st.rerun()
    except Exception:
        try:
            st.experimental_rerun()
        except Exception:
            pass

# ---------------------------------------------------------------------
# Retrieval bridges & simple RAG
# ---------------------------------------------------------------------
# modules/app_helpers.py  (patch the _similarity_search function)

def _similarity_search(store: Any, query: str, k: int) -> List[Any]:
    """
    Call the store's similarity search with broad compatibility.
    """
    if store is None:
        return []

    # 1) Native similarity_search (vectorstores)
    if hasattr(store, "similarity_search"):
        try:
            return store.similarity_search(query, k=k) or []
        except TypeError:
            try:
                return store.similarity_search(query, k) or []
            except Exception:
                return []
        except Exception:
            return []

    # 2) "search" method (some stores expose this)
    if hasattr(store, "search"):
        try:
            return store.search(query, k=k) or []
        except Exception:
            return []

    # 3) LangChain retriever (API changed: .invoke is preferred over .get_relevant_documents)
    if hasattr(store, "as_retriever"):
        try:
            retr = store.as_retriever(search_kwargs={"k": k})
            try:
                # New LC Core: Runnable-like
                docs = retr.invoke(query)                  # ✅ preferred
            except (AttributeError, TypeError):
                # Old API
                docs = retr.get_relevant_documents(query)  # ↩️ fallback
            except Exception:
                # Some retrievers are callable
                try:
                    docs = retr(query)                     # 🧯 last resort
                except Exception:
                    docs = []
            return docs or []
        except Exception:
            return []

    return []

def sp_index_stats(cfg) -> Tuple[int, str]:
    """Return (chunk_count heuristic, index_path) for SharePoint vector store if available."""
    try:
        vs = init_sharepoint_store(cfg)
        path = getattr(vs, "index_path", getattr(vs, "persist_directory", ""))
        docs = _similarity_search(vs, "the", k=1)
        count = 1 if docs else 0
        return (count, str(path) if path else "(in-memory)")
    except Exception:
        return (0, "(unavailable)")

def get_passages_for_query(cfg, query: str, k: int = 6, store: str = "sharepoint") -> List[Dict[str, str]]:
    """Return list of dicts {text, source} from either 'sharepoint' or 'uploaded' store."""
    try:
        vs = init_sharepoint_store(cfg) if store == "sharepoint" else init_uploaded_store(cfg)
    except Exception:
        vs = None
    docs = _similarity_search(vs, query, k=k)
    out: List[Dict[str, str]] = []
    for d in docs:
        text = getattr(d, "page_content", getattr(d, "content", "")) or ""
        meta = getattr(d, "metadata", {}) or {}
        src = meta.get("source") or meta.get("file") or ""
        out.append({"text": text, "source": src})
    return out

def get_sp_docs_any(cfg, query: str, k: int = 6) -> Tuple[List[Any], List[str]]:
    """Convenience for Chat tab: return (docs, sources-list) from SharePoint store."""
    try:
        vs = init_sharepoint_store(cfg)
    except Exception:
        return ([], [])
    docs = _similarity_search(vs, query, k=k)
    sources: List[str] = []
    for d in docs:
        meta = getattr(d, "metadata", {}) or {}
        src = meta.get("source") or meta.get("file") or ""
        if src and src not in sources:
            sources.append(src)
    return (docs, sources)
# modules/app_helpers.py  (drop-in replacement of rag_answer_uploaded)

# --- REPLACE your existing rag_answer_uploaded with this ---
def rag_answer_uploaded(up_store: Any, oai, cfg, prompt: str, top_k: int = 12) -> Tuple[str, List[str]]:
    """
    Cue-driven, lenient RAG: assemble a broad multi-query context first, then ask the model.
    Works for both the big schema pass and sectional passes.
    """
    # 1) Build large, diverse context
    ctx, srcs = _gather_understanding_context(up_store, prompt, base_k=max(8, top_k // 2))

    # 2) If still nothing, token-fallback: query individual tokens to fish some text
    if not ctx.strip():
        q = prompt[:200]
        toks = [t for t in re.sub(r"[/\-]", " ", q.lower()).split() if len(t) >= 4][:10]
        seen = set(); parts = []; srcs = []
        for t in toks:
            try:
                docs = up_store.similarity_search(t, k=4) or []
            except Exception:
                docs = []
            for d in docs:
                txt = getattr(d, "page_content", "") or getattr(d, "content", "") or ""
                meta = getattr(d, "metadata", {}) or {}
                src = meta.get("source", "") or ""
                sig = (src + "|" + txt[:200]).lower()
                if not txt.strip() or sig in seen:
                    continue
                seen.add(sig)
                parts.append(txt[:1800])
                if src and src not in srcs:
                    srcs.append(src)
                if len(parts) >= 24:
                    break
            if len(parts) >= 24:
                break
        ctx = "\n\n---\n\n".join(parts)

    if not ctx.strip():
        return "", []

    # 3) Ask the model tightly
    sys = (
        "You are a precise assistant. Answer ONLY using the provided context. "
        "If something is missing, say 'Not found in provided context'."
    )
    user = f"CONTEXT:\n{ctx}\n\nQUESTION:\n{prompt}\n\nAnswer succinctly. Return JSON exactly when requested."
    try:
        res = oai.chat.completions.create(
            model=getattr(cfg, "ANALYSIS_MODEL", "gpt-4o"),
            messages=[{"role": "system", "content": sys}, {"role": "user", "content": user}],
            temperature=0.0,
            max_tokens= min(2200, getattr(cfg, "MAX_TOKENS", 2200)),
        )
        out = (res.choices[0].message.content or "").strip()
        return out, srcs
    except Exception:
        return "", []

# ---------------------------------------------------------------------
# DOCX helpers: insert Table of Contents at top (used by proposal_builder)
# ---------------------------------------------------------------------
def insert_table_of_contents(doc, title: str = "Table of Contents",
                             heading_style: str = "Heading 1",
                             switches: str = r'TOC \o "1-3" \h \z \u') -> None:
    """
    Insert a Table of Contents field at the very top of the document.

    NOTE: Word will show an empty TOC until fields are updated (Open in Word and press F9).
    """
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except Exception:
        return

    if not getattr(doc, "paragraphs", None):
        doc.add_paragraph("")

    try:
        first_p = doc.paragraphs[0]
        title_p = first_p.insert_paragraph_before(title)
        if heading_style:
            try: title_p.style = heading_style
            except Exception: pass
        toc_p = first_p.insert_paragraph_before("")
    except Exception:
        title_p = doc.add_paragraph(title)
        try: title_p.style = heading_style
        except Exception: pass
        toc_p = doc.add_paragraph("")

    r = toc_p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = switches
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin); r._r.append(instr); r._r.append(fld_sep)
    r2 = toc_p.add_run(); r2._r.append(fld_end)
    doc.add_paragraph("")

# ---------------------------------------------------------------------
# Blueprints (save/load simple generation configs)
# ---------------------------------------------------------------------
def _blueprints_dir(base_dir: Path) -> Path:
    d = base_dir / "blueprints"
    d.mkdir(parents=True, exist_ok=True)
    return d

def list_blueprints(base_dir: Path) -> List[str]:
    d = _blueprints_dir(base_dir)
    return sorted([p.stem for p in d.glob("*.json")])

def save_blueprint(base_dir: Path, name: str, payload: Dict) -> Path:
    d = _blueprints_dir(base_dir)
    path = d / f"{name}.json"
    path.write_text(json.dumps(payload or {}, ensure_ascii=False, indent=2), encoding="utf-8")
    return path

def load_blueprint(base_dir: Path, name: str) -> Optional[Dict]:
    d = _blueprints_dir(base_dir)
    path = d / f"{name}.json"
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None

# ---------------------------------------------------------------------
# Structuring Q&A utilities
# ---------------------------------------------------------------------
def filter_struct_questions(qs: List[Dict], blocked_static: set) -> List[Dict]:
    """
    Remove options that duplicate static-backed sections; drop questions with no remaining options.
    Input `qs` expected like: [{"question": "...", "options": ["Executive Summary", ...]}, ...]
    """
    out: List[Dict] = []
    for q in qs or []:
        opts = [o for o in (q.get("options") or []) if o not in blocked_static]
        if opts:
            out.append({"question": q.get("question", ""), "options": opts})
    return out
