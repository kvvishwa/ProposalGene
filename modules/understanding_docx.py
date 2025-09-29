# modules/understanding_docx.py
from __future__ import annotations

from io import BytesIO
from typing import List, Dict, Any, Optional
from datetime import datetime
import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_XML_ILLEGAL = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')

def _san(s: Optional[str]) -> str:
    if not s:
        return ""
    return _XML_ILLEGAL.sub("", s)

def _add_toc(doc: Document, title: str = "Table of Contents"):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)

def _add_kv(doc: Document, k: str, v: str):
    if not (k or v):
        return
    p = doc.add_paragraph()
    r1 = p.add_run(_san(k or ""))
    r1.bold = True
    if v:
        p.add_run(": ")
        p.add_run(_san(v))

def _add_list(doc: Document, items: List[str], bullet: bool = True):
    for it in items or []:
        if not it:
            continue
        p = doc.add_paragraph(_san(it))
        try:
            p.style = "List Bullet" if bullet else "List Number"
        except Exception:
            pass

def _add_mdish_block(doc: Document, text: str):
    """
    Very light Markdown-to-Word:
      - lines starting with -, *, • → bullet
      - otherwise, plain paragraph
    """
    for raw in (_san(text).splitlines() if text else []):
        ln = raw.strip()
        if not ln:
            doc.add_paragraph("")
            continue
        if ln.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(ln[2:].strip())
            try:
                p.style = "List Bullet"
            except Exception:
                pass
        else:
            doc.add_paragraph(ln)

def _format_date(d: Any) -> str:
    if isinstance(d, dict):
        date = d.get("date") or ""
        time = d.get("time") or ""
        tz = d.get("tz") or ""
        return " ".join(x for x in [date, time, tz] if x)
    return _san(str(d or ""))

def _render_scope(doc: Document, facts: Dict[str, Any]):
    """
    Add a dedicated 'Scope' section. Supports:
      - facts['scope']['in_scope'] and/or ['out_of_scope']
      - facts['scope']['scope_raw'] as fallback
    """
    scope = (facts or {}).get("scope") or {}
    if not scope:
        return
    doc.add_heading("7. Scope", level=1)
    if scope.get("in_scope") or scope.get("out_of_scope"):
        if scope.get("in_scope"):
            doc.add_paragraph("In Scope", style=None)
            _add_list(doc, scope.get("in_scope") or [])
        if scope.get("out_of_scope"):
            doc.add_paragraph("Out of Scope", style=None)
            _add_list(doc, scope.get("out_of_scope") or [])
    elif scope.get("scope_raw"):
        doc.add_paragraph("Scope (Detected)", style=None)
        _add_list(doc, scope.get("scope_raw") or [])

def _render_facts(doc: Document, facts: Dict[str, Any]):
    if not isinstance(facts, dict):
        facts = {}

    # 1) Solicitation
    doc.add_heading("1. Solicitation", level=1)
    sol = facts.get("solicitation", {}) or {}
    _add_kv(doc, "Issuing Entity", sol.get("issuing_entity_name", ""))
    _add_kv(doc, "Entity Type", sol.get("issuing_entity_type", ""))
    _add_kv(doc, "Department/Office", sol.get("department_or_office", ""))
    _add_kv(doc, "Solicitation ID", sol.get("solicitation_id", ""))
    _add_kv(doc, "Title", sol.get("solicitation_title", ""))
    _add_kv(doc, "Communication Policy", sol.get("communication_policy", ""))

    # 2) Points of Contact
    doc.add_heading("2. Points of Contact", level=1)
    for i, poc in enumerate(facts.get("points_of_contact", []) or [], start=1):
        doc.add_heading(f"POC {i}", level=2)
        _add_kv(doc, "Primary", "Yes" if poc.get("primary") else "No")
        _add_kv(doc, "Name", poc.get("name", ""))
        _add_kv(doc, "Title", poc.get("title", ""))
        _add_kv(doc, "Email", poc.get("email", ""))
        _add_kv(doc, "Phone", poc.get("phone", ""))
        _add_kv(doc, "Address", poc.get("address", ""))
        _add_kv(doc, "Notes", poc.get("notes", ""))

    # 3) Schedule
    doc.add_heading("3. Schedule & Deadlines", level=1)
    sch = facts.get("schedule", {}) or {}
    _add_kv(doc, "Release Date", _format_date(sch.get("release_date")))
    pre = sch.get("pre_bid_conference", {}) or {}
    _add_kv(doc, "Pre-bid Conference", f"{_format_date(pre.get('datetime',''))} @ {_san(pre.get('location',''))}")
    site = sch.get("site_visit", {}) or {}
    _add_kv(doc, "Site Visit", f"{_format_date(site.get('datetime',''))} @ {_san(site.get('location',''))}")
    _add_kv(doc, "Q&A Deadline", _format_date(sch.get("qna_deadline")))
    _add_kv(doc, "Addendum Final Date", _format_date(sch.get("addendum_final_date")))
    _add_kv(doc, "Submission Due", _format_date(sch.get("submission_due")))
    _add_kv(doc, "Award Target Date", _format_date(sch.get("award_target_date")))
    _add_kv(doc, "Anticipated Start Date", _format_date(sch.get("anticipated_start_date")))
    _add_kv(doc, "Proposal Validity (days)", str(sch.get("proposal_validity_days") or ""))
    _add_kv(doc, "Contract Term", sch.get("contract_term", ""))
    _add_kv(doc, "Renewal Options", sch.get("renewal_options", ""))

    # 4) Submission Instructions
    doc.add_heading("4. Submission Instructions", level=1)
    sub = facts.get("submission_instructions", {}) or {}
    _add_kv(doc, "Method", sub.get("method", ""))
    _add_kv(doc, "Portal", f"{sub.get('portal_name','')} {sub.get('portal_url','')}".strip())
    _add_kv(doc, "Email", sub.get("email_address", ""))
    _add_kv(doc, "Physical Address", sub.get("physical_address", ""))
    _add_kv(doc, "Copies Required", sub.get("copies_required", ""))
    _add_list(doc, sub.get("format_requirements") or [])
    _add_kv(doc, "Labeling Instructions", sub.get("labeling_instructions", ""))
    _add_kv(doc, "Registration Requirements", sub.get("registration_requirements", ""))
    _add_kv(doc, "Late Submissions Policy", sub.get("late_submissions_policy", ""))

    # 5) Minimum Qualifications
    doc.add_heading("5. Minimum Qualifications", level=1)
    mq = facts.get("minimum_qualifications", {}) or {}
    _add_list(doc, mq.get("licenses_certifications") or [])
    _add_kv(doc, "Years Experience", mq.get("years_experience", ""))
    _add_kv(doc, "Similar Projects", mq.get("similar_projects", ""))
    _add_list(doc, mq.get("insurance_requirements") or [])
    _add_list(doc, mq.get("bonding") or [])
    _add_kv(doc, "Financials", mq.get("financials", ""))
    _add_list(doc, mq.get("other_mandatory") or [])

    # 6) Proposal Organization
    doc.add_heading("6. Proposal Organization", level=1)
    org = facts.get("proposal_organization", {}) or {}
    _add_list(doc, org.get("required_sections") or [])
    for pl in (org.get("page_limits") or []):
        _add_kv(doc, f"Page Limit — {pl.get('section','')}", pl.get("limit",""))
    _add_list(doc, org.get("mandatory_forms") or [])
    _add_kv(doc, "Pricing Format", org.get("pricing_format", ""))
    _add_kv(doc, "Exceptions Policy", org.get("exceptions_policy", ""))

    # 7) Scope (NEW)
    _render_scope(doc, facts)

    # 8) Evaluation & Selection (shifted down by one)
    doc.add_heading("8. Evaluation & Selection", level=1)
    ev = facts.get("evaluation_and_selection", {}) or {}
    for c in ev.get("criteria") or []:
        nm = c.get("name",""); wt = c.get("weight",""); desc = c.get("description","")
        _add_kv(doc, f"Criterion — {nm}", f"{wt} {desc}".strip())
    _add_list(doc, ev.get("pass_fail_criteria") or [])
    oral = ev.get("oral_presentations", {}) or {}
    demo = ev.get("demonstrations", {}) or {}
    _add_kv(doc, "Oral Presentations", f"{'required' if oral.get('required') else 'optional'} — {oral.get('notes','')}".strip())
    _add_kv(doc, "Demonstrations", f"{'required' if demo.get('required') else 'optional'} — {demo.get('notes','')}".strip())
    _add_kv(doc, "BAFO", f"{'allowed' if (ev.get('best_and_final_offers') or {}).get('allowed') else 'not indicated'}")
    _add_kv(doc, "Negotiations", f"{'allowed' if (ev.get('negotiations') or {}).get('allowed') else 'not indicated'}")
    _add_kv(doc, "Selection Summary", ev.get("selection_process_summary",""))
    _add_kv(doc, "Protest Procedure", ev.get("protest_procedure",""))

    # 9) Contract & Compliance (shifted down)
    doc.add_heading("9. Contract & Compliance", level=1)
    cc = facts.get("contract_and_compliance", {}) or {}
    _add_list(doc, cc.get("key_terms") or [])
    _add_list(doc, cc.get("security_privacy") or [])
    if cc.get("sow_summary"):
        doc.add_paragraph(_san(cc.get("sow_summary")))

    # 10) Missing / Notes (shifted down)
    miss = facts.get("missing_notes") or []
    if miss:
        doc.add_heading("10. Missing / Notes", level=1)
        _add_list(doc, miss)

def _render_value_prop(doc: Document, bct_vp_text: Optional[str], bct_vp_sources: Optional[List[str]], generic_vp_text: Optional[str]):
    # 11) BCT Value Proposition (SP-grounded) (shifted down)
    if bct_vp_text:
        doc.add_heading("11. Value Proposition — BCT (SharePoint-grounded)", level=1)
        _add_mdish_block(doc, bct_vp_text)
        if bct_vp_sources:
            doc.add_paragraph("Sources: " + ", ".join(_san(s) for s in bct_vp_sources))

    # 12) Generic Value Proposition (shifted down)
    if generic_vp_text:
        doc.add_heading("12. Value Proposition — Generic", level=1)
        _add_mdish_block(doc, generic_vp_text)

def _render_chat(doc: Document, title: str, messages: List[Dict[str, Any]]):
    doc.add_heading(title, level=1)
    if not messages:
        doc.add_paragraph("— no messages —")
        return
    for turn in messages:
        role = (turn.get("role") or "").strip().lower()
        content = _san(turn.get("content") or "")
        p = doc.add_paragraph()
        r = p.add_run(("User: " if role == "user" else "Assistant: "))
        r.bold = True
        _add_mdish_block(doc, content)
        srcs = turn.get("sources") or []
        if srcs:
            doc.add_paragraph("Sources: " + ", ".join(_san(s) for s in srcs))

def build_understanding_docx(
    facts: Dict[str, Any],
    rfp_chat_messages: List[Dict[str, Any]],
    sp_chat_messages: List[Dict[str, Any]],
    title: str = "RFP Understanding Report",
    add_toc: bool = True,
    *,
    bct_vp_text: Optional[str] = None,
    bct_vp_sources: Optional[List[str]] = None,
    generic_vp_text: Optional[str] = None,
) -> bytes:
    """
    Returns a DOCX as bytes containing:
      - Title page
      - (optional) Table of Contents
      - RFP Facts Summary (with Scope)
      - Value Proposition (BCT grounded + generic)
      - RFP Chat transcript
      - SharePoint Chat transcript
    """
    doc = Document()

    # Title page
    t = doc.add_paragraph(title)
    try:
        t.style = "Title"
    except Exception:
        pass
    ts = doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
    ts.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ToC
    if add_toc:
        _add_toc(doc)

    # Facts (includes Scope)
    _render_facts(doc, facts or {})

    # Value prop sections (after facts)
    _render_value_prop(doc, bct_vp_text, bct_vp_sources, generic_vp_text)

    # Chats
    _render_chat(doc, "RFP Chat Transcript", rfp_chat_messages or [])
    _render_chat(doc, "SharePoint Chat Transcript", sp_chat_messages or [])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
