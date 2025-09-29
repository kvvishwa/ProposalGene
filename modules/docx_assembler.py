# modules/docx_assembler.py
from __future__ import annotations

from copy import deepcopy
from typing import Optional, Tuple

from docx import Document


# ------------------------
# XML helpers (namespace-free)
# ------------------------
def _local_name(tag: str) -> str:
    """Return the localname for an XML tag like '{uri}bookmarkStart' -> 'bookmarkStart'."""
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _iter_all(el):
    """Yield element and all its descendants."""
    yield el
    for ch in getattr(el, "iterchildren", lambda: [])():
        for sub in _iter_all(ch):
            yield sub


def _first_child(el, name: str):
    """First child with the given localname under el (non-recursive)."""
    for ch in getattr(el, "iterchildren", lambda: [])():
        if _local_name(ch.tag) == name:
            return ch
    return None


def _find_ancestor(el, name: str):
    """Climb parents until an ancestor with localname == name is found (or None)."""
    p = getattr(el, "getparent", lambda: None)()
    while p is not None:
        if _local_name(p.tag) == name:
            return p
        p = getattr(p, "getparent", lambda: None)()
    return None


def _normalize(s: str) -> str:
    return "".join(c.lower() for c in (s or "") if c.isalnum())


# ------------------------
# Anchor discovery
# ------------------------
def _find_sdt_anchor(doc: Document, label: str):
    """
    Find a block-level content control (w:sdt) whose alias or tag matches `label` (case/spacing-insensitive).
    Returns the <w:sdt> element if found, else None.
    """
    want = _normalize(label)
    body = doc._element.body

    for el in _iter_all(body):
        if _local_name(el.tag) != "sdt":
            continue

        # sdtPr -> alias/tag (w:alias/@w:val, w:tag/@w:val)
        sdt_pr = _first_child(el, "sdtPr")
        if sdt_pr is None:
            continue

        # Check alias
        alias = _first_child(sdt_pr, "alias")
        if alias is not None:
            val = alias.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
            if _normalize(val) == want:
                return el

        # Check tag
        tag = _first_child(sdt_pr, "tag")
        if tag is not None:
            val = tag.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
            if _normalize(val) == want:
                return el

    return None


def _find_bookmark_anchor(doc: Document, label: str):
    """
    Find a bookmark start (w:bookmarkStart) whose @w:name matches `label` (case/spacing-insensitive).
    Returns the paragraph <w:p> containing it (so we can insert after), else None.
    """
    want = _normalize(label)
    body = doc._element.body

    for el in _iter_all(body):
        if _local_name(el.tag) != "bookmarkStart":
            continue
        name = el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name")
        if _normalize(name) == want:
            # Insert after the paragraph that contains the bookmark
            host_p = _find_ancestor(el, "p")
            return host_p
    return None


def _find_marker_paragraph(doc: Document, label: str):
    """
    Find a paragraph containing a simple text marker like [[Label]] (case-insensitive).
    Returns the paragraph element <w:p> or None.
    """
    want = f"[[{label}]]".lower()
    for p in doc.paragraphs:
        if want in (p.text or "").lower():
            return p._p
    return None


# ------------------------
# Insertion primitives
# ------------------------
def _insert_blocks_into_sdt(sdt_el, frag_doc: Document):
    """
    Replace the content of a block-level SDT (w:sdt) with the blocks from frag_doc.
    """
    sdt_content = _first_child(sdt_el, "sdtContent")
    if sdt_content is None:
        # malformed SDT; nothing we can do
        return False

    # Clear existing content
    for ch in list(getattr(sdt_content, "iterchildren", lambda: [])()):
        sdt_content.remove(ch)

    # Append deep-copied blocks from fragment's body
    for blk in frag_doc._element.body.iterchildren():
        sdt_content.append(deepcopy(blk))
    return True


def _insert_blocks_after(host_block_el, frag_doc: Document):
    """
    Insert blocks from frag_doc after the given host block element (e.g., paragraph <w:p>).
    """
    parent = getattr(host_block_el, "getparent", lambda: None)()
    if parent is None:
        return False

    after = host_block_el
    for blk in frag_doc._element.body.iterchildren():
        new_blk = deepcopy(blk)
        parent.insert(parent.index(after) + 1, new_blk)
        after = new_blk
    return True


def _append_blocks_to_document(doc: Document, frag_doc: Document, heading_text: Optional[str] = None):
    """
    Append an optional Heading 1 and all blocks from frag_doc to end of doc.
    """
    if heading_text:
        h = doc.add_paragraph(heading_text)
        h.style = "Heading 1"
    # Append blocks at XML level for full fidelity
    body = doc._element.body
    for blk in frag_doc._element.body.iterchildren():
        body.append(deepcopy(blk))
    return True


# ------------------------
# Public API
# ------------------------
def insert_section(
    base_doc: Document,
    section_label: str,
    section_doc: Document,
    try_anchors: bool = True,
    add_heading_if_appending: bool = True,
) -> Tuple[bool, str]:
    """
    Insert `section_doc` into `base_doc`:
      1) If try_anchors: try block SDT with alias/tag == section_label,
         else bookmark [[section_label]] or w:bookmarkStart name==label,
         else append at end.
      2) When appending, optionally prepend a Heading 1 with section_label.

    Returns (success, where) where 'where' ∈ {'sdt','bookmark','marker','append'}.
    """
    # 1) Try content-control anchor (block-level SDT)
    if try_anchors:
        sdt = _find_sdt_anchor(base_doc, section_label)
        if sdt is not None:
            ok = _insert_blocks_into_sdt(sdt, section_doc)
            return ok, "sdt"

        # 2) Try bookmark anchor (named bookmark)
        host_p = _find_bookmark_anchor(base_doc, section_label)
        if host_p is not None:
            ok = _insert_blocks_after(host_p, section_doc)
            return ok, "bookmark"

        # 3) Try simple text marker [[Label]]
        marker_p = _find_marker_paragraph(base_doc, section_label)
        if marker_p is not None:
            # Remove the marker paragraph and insert content at that spot
            parent = marker_p.getparent()
            idx = parent.index(marker_p)
            parent.remove(marker_p)
            after = parent[idx - 1] if idx > 0 else None
            # If there's a previous block, insert after it; else insert at start
            if after is not None:
                for blk in section_doc._element.body.iterchildren():
                    new_blk = deepcopy(blk)
                    parent.insert(parent.index(after) + 1, new_blk)
                    after = new_blk
            else:
                # insert at the beginning of body
                for blk in reversed(list(section_doc._element.body.iterchildren())):
                    parent.insert(0, deepcopy(blk))
            return True, "marker"

    # 4) Fallback: append to end (optionally with Heading 1)
    _append_blocks_to_document(
        base_doc,
        section_doc,
        heading_text=(section_label if add_heading_if_appending else None),
    )
    return True, "append"
