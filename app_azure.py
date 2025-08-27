# app_azure.py
from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path
from copy import deepcopy
from typing import List, Tuple

import gradio as gr
from fastapi import FastAPI
from openai import OpenAI
import docx

# your config & modules
from config import load_config, Config
from modules.docx_assembler import insert_section
from modules.docx_generator import generate_response_docx_from_text
from modules.text_extraction import extract_text

from modules.vectorstore import (
    init_uploaded_store, ingest_files, init_sharepoint_store
)
from modules.utils import save_to_temp, cleanup_temp_files

from modules.app_helpers import (
    # FS / mapping
    ensure_dirs, list_files, load_static_map,
    # SP + RAG
    get_sp_docs_any, rag_answer_uploaded, sp_index_stats,
    # recos generator (SharePoint grounded)
    generate_dynamic_recommendations,
)

# ---------------------------
# App-level config & state
# ---------------------------
cfg = load_config()
oai = OpenAI(api_key=cfg.OPENAI_API_KEY)

BASE_DIR = Path(os.getenv("BASE_DIR", getattr(cfg, "BASE_DIR", "/home/site/data")))
VECTORSTORE_DIR = Path(os.getenv("VECTORSTORE_DIR", getattr(cfg, "VECTORSTORE_DIR", "/home/site/vectorstore")))
STATIC_DIR = BASE_DIR / "static_sources"
TEMPLATE_DIR = BASE_DIR / "templates"
STATIC_MAP_FILE = BASE_DIR / "static_map.json"

ensure_dirs(BASE_DIR)
STATIC_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

STATIC_SECTIONS = [
    "Cover Letter",
    "Executive Summary",
    "Experience",
    "Offerings",
    "References",
    "Team & Credentials",
    "Case Studies",
]

STATE = {
    "up_store": None,          # vectorstore for uploaded RFP docs
    "uploaded_tempfiles": [],  # temp paths to clean up
    "sp_chat": [],             # [(user, assistant), ...]
    "rfp_chat": [],            # [(user, assistant), ...]
}

# ---------------------------
# Gradio logic functions
# ---------------------------

def sp_index_status() -> str:
    chunks, vec_path = sp_index_stats(cfg)
    return f"SharePoint index path: `{vec_path}` • Chunks: **{chunks}**"


def sp_chat_fn(message: str, history: List[Tuple[str, str]]) -> Tuple[List[Tuple[str, str]], str]:
    """
    Chat over SharePoint vector DB (no category restriction).
    Returns updated history and a 'Sources: ...' line for the last answer.
    """
    if not message or not message.strip():
        return history, ""
    q = message.strip()
    docs, sources = get_sp_docs_any(cfg, q, k=6)
    if not docs:
        ans = "No relevant passages found in the SharePoint index. Try Pull & Index (Settings) or broaden your query."
        history = history + [(q, ans)]
        return history, ""
    context = "\n\n".join(d.page_content for d in docs)
    prompt = f"""Answer the question strictly using this SharePoint context.

CONTEXT:
{context}

QUESTION:
{q}
"""
    try:
        res = oai.chat.completions.create(
            model=cfg.ANALYSIS_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        ans = res.choices[0].message.content.strip()
    except Exception as ex:
        ans = f"LLM error: {ex}"
    history = history + [(q, ans)]
    src_line = "Sources: " + ", ".join(sources) if sources else ""
    return history, src_line


def rfp_upload_fn(files: List[gr.File]) -> str:
    """
    Upload and vectorize RFP documents into a session-level store.
    """
    if not files:
        return "No files selected."
    # init or reuse store
    if STATE["up_store"] is None:
        STATE["up_store"] = init_uploaded_store(cfg)

    n = 0
    for f in files:
        # gr.File provides a path
        path = Path(f.name)
        if not path.exists():
            # Fallback in case of different object type
            continue
        ingest_files([str(path)], STATE["up_store"], getattr(cfg, "CHUNK_SIZE", 1000))
        STATE["uploaded_tempfiles"].append(str(path))
        n += 1
    return f"Ingested and indexed {n} file(s) for RFP understanding."


def rfp_chat_fn(message: str, history: List[Tuple[str, str]]) -> Tuple[List[Tuple[str, str]], str]:
    if not message or not message.strip():
        return history, ""
    q = message.strip()
    if STATE["up_store"] is None:
        ans = "Please upload and vectorize RFP documents first."
        return history + [(q, ans)], ""

    ans, sources = rag_answer_uploaded(STATE["up_store"], oai, cfg, q, top_k=6)
    history = history + [(q, ans)]
    src_line = "Sources: " + ", ".join(sources) if sources else ""
    return history, src_line


def generate_draft_and_recos_fn(
    template_name: str,
    static_selected: List[str],
    dynamic_selected: List[str],
    include_sources: bool,
    context_topk: int,
    tone: str,
) -> Tuple[str, str, str]:
    """
    Produce two DOCX files:
      1) Draft response: Template + selected static sections + dynamic placeholders with SharePoint-grounded recommendations.
      2) Recommendations-only DOCX.
    Returns file paths and a status line.
    """
    if not template_name or template_name == "— choose a template —":
        return "", "", "Please choose a template."

    template_path = TEMPLATE_DIR / template_name
    if not template_path.exists():
        return "", "", f"Template not found: {template_name}"

    # load map and check files
    mapping = load_static_map(STATIC_MAP_FILE)
    static_paths = {sec: str((STATIC_DIR / mapping[sec]).resolve()) for sec in static_selected if mapping.get(sec)}

    # Build main Draft doc
    base_doc = docx.Document(str(template_path))

    # Also build Recommendations-only doc
    recs_doc = docx.Document()
    recs_doc.add_heading("Dynamic Recommendations", level=1)

    # Insert static sections
    for label in static_selected:
        fpath = static_paths.get(label)
        if not fpath or not os.path.exists(fpath):
            continue
        frag_doc = docx.Document(fpath)
        _ = insert_section(
            base_doc,
            section_label=label,
            section_doc=frag_doc,
            try_anchors=True,
            add_heading_if_appending=True,
        )

    # Insert dynamics with SharePoint-grounded recos
    for label in dynamic_selected:
        rec_md, rec_srcs = generate_dynamic_recommendations(
            oai, cfg, label,
            k_ctx=int(context_topk),
            max_bullets=8,
            tone=tone or "Professional",
        )
        # Draft: put a heading + short intro + bullets
        draft_md = f"## {label}\n\n> _Placeholder_: refine this section using the recommendations below.\n\n{rec_md}\n"
        frag_bytes = generate_response_docx_from_text(draft_md)
        frag_doc = docx.Document(frag_bytes)
        _ = insert_section(
            base_doc,
            section_label=label,
            section_doc=frag_doc,
            try_anchors=True,
            add_heading_if_appending=True,
        )
        if include_sources and rec_srcs:
            p = base_doc.add_paragraph()
            run = p.add_run("Sources: " + ", ".join(rec_srcs))
            run.italic = True

        # Recs-only doc
        recs_doc.add_heading(label, level=2)
        rec_frag = docx.Document(generate_response_docx_from_text(rec_md))
        for blk in rec_frag._element.body.iterchildren():
            recs_doc._element.body.append(deepcopy(blk))
        if include_sources and rec_srcs:
            p = recs_doc.add_paragraph()
            r = p.add_run("Sources: " + ", ".join(rec_srcs))
            r.italic = True

    # Save under /home/site if present, else /tmp
    out_dir = Path("/home/site/wwwroot") if Path("/home/site/wwwroot").exists() else Path("/tmp")
    out_dir.mkdir(parents=True, exist_ok=True)
    draft_path = out_dir / "Draft_Response_Static+Dynamic.docx"
    recs_path  = out_dir / "Dynamic_Recommendations.docx"

    bio1, bio2 = BytesIO(), BytesIO()
    base_doc.save(bio1); bio1.seek(0)
    recs_doc.save(bio2); bio2.seek(0)
    draft_path.write_bytes(bio1.getvalue())
    recs_path.write_bytes(bio2.getvalue())

    return str(draft_path), str(recs_path), "Draft and recommendations generated."

# ---------------------------
# Gradio UI (Tabs)
# ---------------------------
def gradio_ui() -> gr.Blocks:
    templates = ["— choose a template —"] + [f for f in os.listdir(TEMPLATE_DIR) if f.lower().endswith(".docx")]
    mapping = load_static_map(STATIC_MAP_FILE)
    available_static = [sec for sec in STATIC_SECTIONS if mapping.get(sec)]

    dynamic_pool = [
        "Executive Summary",
        "Approach & Methodology",
        "Staffing Plan & Roles",
        "Transition & Knowledge Transfer",
        "Service Levels & Governance",
        "Assumptions & Exclusions",
        "Risk & Mitigation Plan",
        "Project Governance",
        "Change Management",
        "Quality Management",
        "Compliance & Security",
        "Timeline & Milestones",
    ]

    with gr.Blocks(title="Proposal Studio") as demo:
        gr.Markdown("## 📄 Proposal Studio (FastAPI + Gradio)")

        with gr.Tab("1) Chat (SharePoint DB)"):
            idx = gr.Markdown(sp_index_status())
            chatbot = gr.Chatbot(type="messages", height=420)
            sp_msg = gr.Textbox(label="Ask anything from your SharePoint knowledge base…", lines=3)
            sp_src = gr.Markdown()
            with gr.Row():
                sp_send = gr.Button("Send", variant="primary")
                sp_clear = gr.Button("Clear chat")

            def _sp_send(m, hist):
                hist = hist or []
                new_hist, src = sp_chat_fn(m, hist)
                return new_hist, gr.update(value=""), src

            sp_send.click(_sp_send, inputs=[sp_msg, chatbot], outputs=[chatbot, sp_msg, sp_src])
            sp_msg.submit(_sp_send, inputs=[sp_msg, chatbot], outputs=[chatbot, sp_msg, sp_src])

            def _sp_clear():
                return [], ""
            sp_clear.click(_sp_clear, outputs=[chatbot, sp_src])

        with gr.Tab("2) Proposal Understanding"):
            gr.Markdown("Upload RFP docs and chat with them (RAG over uploads).")
            up = gr.File(label="Upload multiple files", file_count="multiple", file_types=[".pdf", ".docx", ".pptx"])
            up_btn = gr.Button("Vectorize")
            up_status = gr.Markdown()

            up_btn.click(rfp_upload_fn, inputs=[up], outputs=[up_status])

            rfp_chat = gr.Chatbot(type="messages", height=420)
            rfp_msg = gr.Textbox(label="Ask about the uploaded documents…", lines=3)
            rfp_src = gr.Markdown()
            with gr.Row():
                rfp_send = gr.Button("Send", variant="primary")
                rfp_clear = gr.Button("Clear chat")

            def _rfp_send(m, hist):
                hist = hist or []
                new_hist, src = rfp_chat_fn(m, hist)
                return new_hist, gr.update(value=""), src

            rfp_send.click(_rfp_send, inputs=[rfp_msg, rfp_chat], outputs=[rfp_chat, rfp_msg, rfp_src])
            rfp_msg.submit(_rfp_send, inputs=[rfp_msg, rfp_chat], outputs=[rfp_chat, rfp_msg, rfp_src])
            rfp_clear.click(lambda: ([], ""), outputs=[rfp_chat, rfp_src])

        with gr.Tab("3) Draft Response (Template + Static + Dynamic recos)"):
            with gr.Row():
                tpl = gr.Dropdown(templates, value=templates[0], label="Template (.docx) from /templates")
                include_sources = gr.Checkbox(True, label="Append 'Sources' lines for dynamic sections")
            with gr.Row():
                static_sel = gr.CheckboxGroup(choices=available_static, value=available_static, label="Static sections (from mapped .docx)")
            with gr.Row():
                dyn_sel = gr.CheckboxGroup(
                    choices=dynamic_pool,
                    value=[x for x in ["Executive Summary", "Approach & Methodology", "Staffing Plan & Roles"] if x in dynamic_pool],
                    label="Dynamic sections (placeholders + SharePoint-grounded recommendations)"
                )
            with gr.Row():
                topk = gr.Slider(3, 12, value=6, step=1, label="Context Top-K (SharePoint)")
                tone = gr.Dropdown(choices=["Professional","Persuasive","Technical","Executive"], value="Professional", label="Tone")

            go = gr.Button("Generate Draft + Recommendations", variant="primary")
            draft_out = gr.File(label="Draft DOCX")
            recs_out = gr.File(label="Recommendations DOCX")
            status = gr.Markdown()

            def _gen(tpl_name, statics, dyns, inc_src, k, tn):
                draft, recs, msg = generate_draft_and_recos_fn(tpl_name, statics or [], dyns or [], bool(inc_src), int(k), tn)
                return gr.update(value=draft if draft else None), gr.update(value=recs if recs else None), msg

            go.click(_gen, inputs=[tpl, static_sel, dyn_sel, include_sources, topk, tone], outputs=[draft_out, recs_out, status])

    return demo


# ---------------------------
# FastAPI app + mount Gradio
# ---------------------------
fastapi_app = FastAPI(title="Proposal Studio API")
demo = gradio_ui()
app = gr.mount_gradio_app(fastapi_app, demo, path="/")

# Optional: add a simple health endpoint
@fastapi_app.get("/healthz")
def healthz():
    return {"ok": True}
